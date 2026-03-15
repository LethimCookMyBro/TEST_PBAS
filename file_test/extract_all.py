import PyPDF2
import docx
import os
import sys

folder = r"c:\Users\User\Downloads\PBAS\file_test"
output_folder = r"c:\Users\User\Downloads\PBAS\file_test\extracted"
os.makedirs(output_folder, exist_ok=True)

# Extract PDFs
pdf_files = [f for f in os.listdir(folder) if f.endswith('.pdf')]
for pdf_file in pdf_files:
    print(f"--- Extracting: {pdf_file} ---")
    try:
        reader = PyPDF2.PdfReader(os.path.join(folder, pdf_file))
        text = ""
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n=== Page {i+1} ===\n{page_text}"
        out_name = pdf_file.replace('.pdf', '.txt')
        with open(os.path.join(output_folder, out_name), 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"  Extracted {len(reader.pages)} pages, {len(text)} chars")
    except Exception as e:
        print(f"  ERROR: {e}")

# Extract DOCX
docx_files = [f for f in os.listdir(folder) if f.endswith('.docx')]
for docx_file in docx_files:
    print(f"--- Extracting: {docx_file} ---")
    try:
        doc = docx.Document(os.path.join(folder, docx_file))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # also get tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                text += row_text + "\n"
        out_name = docx_file.replace('.docx', '.txt')
        with open(os.path.join(output_folder, out_name), 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"  Extracted {len(doc.paragraphs)} paragraphs, {len(text)} chars")
    except Exception as e:
        print(f"  ERROR: {e}")

print("\nDone!")
